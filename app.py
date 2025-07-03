from flask import Flask, render_template, request, jsonify, send_file
from io import BytesIO
from docx import Document
import logging
import json
import requests
import matplotlib.pyplot as plt
from datetime import datetime
import os
import platform
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, Any

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB upload limit

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# ======================
# Ollama Configuration
# ======================
OLLAMA_CONFIG = {
    "base_url": "http://localhost:11434/api",
    "model": "phi3:mini",
    "options": {
        "num_ctx": 1024,
        "num_thread": max(1, os.cpu_count() - 1),
        "temperature": 0.3,
        "stop": ["\n###", "\n##"]
    }
}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        logger.debug("Received analyze request")

        if not request.is_json:
            return jsonify({'error': 'Request must be JSON'}), 400

        audit_data = request.get_json()
        if not audit_data or 'audit_report' not in audit_data:
            return jsonify({'error': "Data must contain 'audit_report' object"}), 400

        audit_report = audit_data['audit_report']
        exceptions = audit_report.get('exceptions', [])
        if not exceptions:
            return jsonify({'error': "'audit_report' must contain 'exceptions' array"}), 400

        logger.info(f"Processing {len(exceptions)} exceptions for {audit_report.get('branch', 'unknown branch')}")

        try:
            prompt_batches = prepare_audit_prompt(audit_data)
        except Exception as e:
            return jsonify({'error': f"Failed to prepare analysis prompt: {str(e)}"}), 400

        try:
            analysis_result = query_ollama(prompt_batches)
            if 'error' in analysis_result:
                return jsonify(analysis_result), 502
        except requests.exceptions.Timeout:
            return jsonify({'error': 'Analysis timeout - try a smaller dataset'}), 504

        try:
            chart_urls = generate_charts(audit_report)
        except Exception as e:
            logger.error(f"Chart generation failed: {str(e)}")
            chart_urls = {}

        response_data = {
            'analysis': analysis_result,
            'charts': chart_urls,
            'metadata': {
                'branch': audit_report.get('branch'),
                'period': audit_report.get('period'),
                'exception_count': len(exceptions)
            }
        }

        return jsonify(response_data)

    except json.JSONDecodeError as e:
        return jsonify({'error': 'Invalid JSON format'}), 400
    except Exception as e:
        logger.exception("Unexpected analysis error")
        return jsonify({'error': f'Internal server error: {str(e)}'}), 500

@app.route('/generate-report', methods=['POST'])
def generate_report():
    try:
        data = request.json
        if not data or 'analysis' not in data:
            return jsonify({'error': 'Invalid report data'}), 400

        doc = generate_word_document(data['analysis'])

        return send_file(
            doc,
            as_attachment=True,
            download_name='Audit_Report.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f"Report generation failed: {str(e)}"}), 500

# ======================
# Core Functions
# ======================
def prepare_audit_prompt(audit_data, batch_size=500):
    branch = audit_data['audit_report'].get('branch', 'Unknown branch')
    period = audit_data['audit_report'].get('period', 'Unknown period')
    exceptions = audit_data['audit_report'].get('exceptions', [])

    prompts = []
    for i in range(0, len(exceptions), batch_size):
        batch = exceptions[i:i+batch_size]
        prompt = f"""You are an expert banking auditor analyzing exceptions from {branch} branch for {period}.
Batch {i//batch_size + 1}:
Provide:
1. Executive summary
2. Findings (title, description, risk_level, impact)
3. Recommendations

Exceptions:
{json.dumps(batch, indent=2)}"""
        prompts.append(prompt)
    return prompts

def query_ollama(prompt_batches):
    results = []

    def call_ollama(prompt):
        try:
            payload = {
                "model": OLLAMA_CONFIG["model"],
                "prompt": prompt,
                "stream": False,
                "format": "json",
                "options": OLLAMA_CONFIG["options"]
            }
            response = requests.post(
                f"{OLLAMA_CONFIG['base_url']}/generate",
                json=payload,
                timeout=90
            )
            response.raise_for_status()
            return json.loads(response.text)
        except Exception as e:
            logger.error(f"Error in Ollama batch: {e}")
            return {"error": str(e)}

    with ThreadPoolExecutor(max_workers=4) as executor:
        future_to_prompt = {executor.submit(call_ollama, p): p for p in prompt_batches}
        for future in as_completed(future_to_prompt):
            results.append(future.result())

    combined = {
        "summary": " ".join(r.get("summary", "") for r in results if isinstance(r, dict)),
        "findings": sum((r.get("findings", []) for r in results if isinstance(r, dict)), []),
        "participants": sum((r.get("participants", []) for r in results if isinstance(r, dict)), []),
        "trends": " ".join(r.get("trends", "") for r in results if isinstance(r, dict))
    }
    return combined

def generate_charts(audit_data: Dict[str, Any]) -> Dict[str, str]:
    temp_dir = os.path.join(os.environ.get('TEMP', ''), 'audit_charts')
    os.makedirs(temp_dir, exist_ok=True)

    severities = count_severities(audit_data)
    plt.figure(figsize=(8, 4))
    plt.bar(severities.keys(), severities.values())
    plt.title('Exception Severity Distribution')
    severity_path = os.path.join(temp_dir, 'severity.png')
    plt.savefig(severity_path)
    plt.close()

    return {'severity': severity_path.replace('\\', '/')}

def generate_word_document(analysis: Dict[str, Any]) -> BytesIO:
    doc = Document()
    doc.add_heading('Banking Audit Exception Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(analysis.get('summary', 'No summary provided'))

    doc.add_heading('Detailed Findings', level=1)
    for finding in analysis.get('findings', []):
        doc.add_heading(finding.get('title', 'Untitled Finding'), level=2)
        doc.add_paragraph(finding.get('description', 'No description'))
        doc.add_paragraph(f"Severity: {finding.get('severity', 'unknown').upper()}")
        doc.add_paragraph(f"Recommendation: {finding.get('recommendation', 'None')}")

    if 'participants' in analysis:
        doc.add_heading('Participants Involved', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'LightShading-Accent1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Name'
        hdr[1].text = 'Role'
        hdr[2].text = 'Branch'

        for p in analysis['participants']:
            row = table.add_row().cells
            row[0].text = p.get('name', '')
            row[1].text = p.get('role', '')
            row[2].text = p.get('branch', '')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def count_severities(data: Dict[str, Any]) -> Dict[str, int]:
    counts = {'low': 0, 'medium': 0, 'high': 0}
    for item in data.get('exceptions', []):
        sev = item.get('severity', '').lower()
        if sev in counts:
            counts[sev] += 1
    return counts

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True, threaded=True)
